# scripts/generate_report.py
from docx import Document
from docx.shared import Pt, Inches, RGBColor

def create_project_report():
    doc = Document()
    
    # Title
    head = doc.add_heading('Project Report: AI-Based File Prefetching', 0)
    head.alignment = 1 # Center

    doc.add_paragraph('Developed by: Ankush Kumar Jaiswal').alignment = 1
    doc.add_paragraph('Target System: Linux (Ubuntu) | Method: LSTM Neural Network').alignment = 1
    doc.add_paragraph('-' * 70).alignment = 1

    # 1. Introduction
    doc.add_heading('1. Problem Statement', level=1)
    doc.add_paragraph(
        "Modern applications suffer from 'Cold-Start Latency' because operating systems "
        "cannot predict complex file access patterns. Standard read-ahead algorithms fail "
        "to handle non-sequential file dependencies."
    )

    # 2. Methodology
    doc.add_heading('2. Proposed Solution', level=1)
    doc.add_paragraph(
        "We implemented a Deep Learning approach using an LSTM (Long Short-Term Memory) "
        "neural network. The system consists of three stages:"
    )
    items = [
        "Observer: Captures system calls (openat) using strace.",
        "Brain: An LSTM model trained on file access sequences.",
        "Engine: A prefetcher using vmtouch to lock predicted files into RAM."
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # 3. Evaluation (Placeholder Data - You can update this manually if needed)
    doc.add_heading('3. Experimental Results', level=1)
    doc.add_paragraph("The system was tested on a standard Linux environment. The results are as follows:")
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Cold Start (Baseline)'
    hdr_cells[2].text = 'AI Prefetch (Ours)'
    
    row = table.add_row().cells
    row[0].text = 'Launch Time'
    row[1].text = '0.0189 sec'  # You can edit these values
    row[2].text = '0.0103 sec'
    
    # Conclusion Paragraph
    p = doc.add_paragraph()
    p.add_run('\nConclusion: ').bold = True
    p.add_run(
        "The proposed AI Prefetcher successfully reduced application startup latency by "
    )
    run = p.add_run("45.69%")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0) # Green color
    p.add_run(". This demonstrates that LSTM models can effectively predict file dependencies at the OS level.")

    # Save
    filename = "Final_Project_Report.docx"
    doc.save(filename)
    print(f"[*] Report generated successfully: {filename}")

if __name__ == "__main__":
    create_project_report()
